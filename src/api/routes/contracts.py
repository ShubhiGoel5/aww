"""
Contract Management - Multi-tenant CRUD + AI Review
Upload, manage, and review contracts with Claude AI
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Query, Request
from typing import Optional, List, Dict
from pydantic import BaseModel
from datetime import datetime, date, timedelta
import os
import json
import uuid
from pathlib import Path
import psycopg2
from psycopg2.extras import RealDictCursor
from ..middleware.auth import get_current_user, get_db
import httpx
import sys

# FIX 3, FIX 5: Import security utilities
sys.path.insert(0, str(Path(__file__).parent.parent))
from security_utils import (
    sanitize_filename, validate_file_path, check_content_length,
    MAX_FILE_SIZE, ALLOWED_EXTENSIONS, sanitize_log
)

router = APIRouter(prefix="/v1/contracts", tags=["contracts"])

# File upload config
# Use /tmp which is always writable on cloud hosts like Render
UPLOAD_DIR = Path("/tmp/legal-ai-agent-uploads/contracts")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def extract_file_text(file_path: str, file_ext: str, content: bytes = None) -> Optional[str]:
    """Extract text from PDF, DOCX, or TXT files"""
    import io
    try:
        if file_ext == ".txt":
            return content.decode('utf-8', errors='ignore') if content else ""
        
        elif file_ext == ".pdf":
            from pypdf import PdfReader
            # Read from file path or from bytes in memory
            source = file_path if file_path else io.BytesIO(content)
            reader = PdfReader(source)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts) if text_parts else None
        
        elif file_ext in (".docx", ".doc"):
            import unicodedata
            from docx import Document
            source = file_path if file_path else io.BytesIO(content)
            doc = Document(source)
            text_parts = []
            for para in doc.paragraphs:
                text = unicodedata.normalize('NFC', para.text.strip())
                if text:
                    if para.style and para.style.name and 'heading' in para.style.name.lower():
                        text_parts.append('**' + text + '**')
                    else:
                        text_parts.append(text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(unicodedata.normalize('NFC', cell.text.strip()) for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n".join(text_parts) if text_parts else None
        
        elif file_ext in (".jpg", ".jpeg", ".png"):
            return None
    except Exception as e:
        print(f"Text extraction error ({file_ext}): {e}")
        return None


async def ai_analyze_contract(text: str, company_id: Optional[str] = None) -> dict:
    """Use LLM to extract contract metadata from text"""
    from src.services.llm_provider import call_llm_simple
    
    system_prompt = """Analyze the contract text and extract information. Return pure JSON (no markdown):
{
    "title": "Contract name (e.g. Employment Contract between ABC Corp and John Doe)",
    "contract_type": "employment_contract|service_contract|lease_contract|sales_contract|partnership_contract|other",
    "parties": ["Party A: ABC Corp", "Party B: John Doe"],
    "start_date": "YYYY-MM-DD or null",
    "end_date": "YYYY-MM-DD or null",
    "value": amount (integer) or null,
    "summary": "Summary of main content in 2-3 sentences"
}
Return ONLY JSON, no text or markdown."""

    try:
        result = await call_llm_simple(system_prompt, f"Analyze contract:\n\n{text[:15000]}", max_tokens=1024, company_id=company_id)
        content = result["content"].strip()
        # Clean markdown code blocks if any
        if content.startswith("```"):
            content = content.split("\n", 1)[1] if "\n" in content else content
            content = content.rsplit("```", 1)[0] if "```" in content else content
        return json.loads(content.strip())
    except Exception as e:
        print(f"AI contract analysis error: {e}")
        return {}


# ============================================
# Models
# ============================================

class ContractCreate(BaseModel):
    name: str
    contract_type: Optional[str] = None
    parties: Optional[List[Dict]] = []
    start_date: Optional[date] = None
    end_date: Optional[date] = None
    notes: Optional[str] = None
    metadata: Optional[Dict] = {}

class ContractUpdate(BaseModel):
    name: Optional[str] = None
    contract_type: Optional[str] = None
    parties: Optional[List[Dict]] = None
    start_date: Optional[date] = None
    end_date: Optional[date] = None
    status: Optional[str] = None
    notes: Optional[str] = None
    metadata: Optional[Dict] = None

class ContractReview(BaseModel):
    contract_id: Optional[str] = None
    contract_text: Optional[str] = None
    focus_areas: Optional[List[str]] = []

# ============================================
# Helpers
# ============================================

async def call_llm_for_review(contract_text: str, contract_type: str = None, company_id: Optional[str] = None) -> dict:
    """Call LLM for contract review"""
    from src.services.llm_provider import call_llm_simple
    
    system_prompt = """You are a senior lawyer specializing in contract review under Indian law, specifically the Industrial Relations Code, 2020 (IR Code), Indian Contract Act 1872, Companies Act 2013, and related guiding documents.

TASK: Analyze EACH clause in the contract in DETAIL according to the following criteria:

1. **CLAUSE-BY-CLAUSE ANALYSIS**: Read and evaluate each specific clause. Identify which clauses are favorable/unfavorable for each party.

2. **LEGAL COMPLIANCE CHECK**: 
   - Compare with IR Code (if it's an employment contract): probation period, minimum wage, working hours, leave, social insurance, unilateral termination, dismissal.
   - Compare with Indian Contract Act 1872: contract validity, void contracts, penalties, damages.
   - Relevant specialized laws.

3. **MANDATORY CLAUSES CHECK**: Identify mandatory clauses required by law that the contract is missing.

4. **RISK ASSESSMENT**: Analyze legal risks for each party, including dispute risks, invalidity risks, and administrative penalty risks.

5. **SUGGESTED REVISIONS**: Specifically propose revisions/additions with legal basis.

Return as pure JSON (no markdown):
{
    "risk_score": <number from 1-100, where: 1-30=low, 31-60=medium, 61-100=high>,
    "risk_level": "low|medium|high",
    "issues": [
        {
            "type": "violation|missing|risk|unfavorable|suggestion",
            "severity": "critical|high|medium|low",
            "clause": "Name/number of relevant clause",
            "description": "Detailed description of the issue",
            "affected_party": "Party A|Party B|Both parties",
            "legal_basis": "Specific legal basis (section, sub-section, act)",
            "recommendation": "Specific revision proposal"
        }
    ],
    "recommendations": ["Overall recommendation 1", "Overall recommendation 2"],
    "summary": "3-5 paragraphs summarizing the overall contract analysis, including: (1) general comments on structure and content, (2) main issues found, (3) assessment of the balance of rights between the parties, (4) main legal risks, (5) overall recommendations.",
    "overall_assessment": "Detailed overall assessment of the legal quality of the contract, compliance with the law, and ability to protect the parties' rights."
}

IMPORTANT:
- Must analyze EACH clause SPECIFICALLY, no generic evaluations
- Summary must be 3-5 detailed paragraphs, NOT just 1 sentence
- Each issue must have a complete legal_basis and recommendation
- Must check for mandatory clauses under the law
- Return ONLY pure JSON, no text or markdown"""

    user_message = f"""CONTRACT FOR REVIEW:
{contract_text[:30000]}

{f"CONTRACT TYPE: {contract_type}" if contract_type else ""}

Please review and return the JSON as requested."""

    result = await call_llm_simple(system_prompt, user_message, max_tokens=8192, company_id=company_id)
    
    content = result["content"]
    
    # Try to parse JSON from LLM's response
    try:
        # Clean markdown if present
        content_clean = content.strip()
        if content_clean.startswith("```"):
            lines = content_clean.split("\n")
            json_lines = []
            in_block = False
            for line in lines:
                if line.startswith("```") and not in_block:
                    in_block = True
                    continue
                elif line.startswith("```") and in_block:
                    break
                elif in_block:
                    json_lines.append(line)
            content_clean = "\n".join(json_lines)
        review_result = json.loads(content_clean)
    except:
        # If not valid JSON, wrap it
        review_result = {"raw_analysis": content}
    
    return {
        "review": review_result,
        "tokens": {
            "input": result.get("input_tokens", 0),
            "output": result.get("output_tokens", 0)
        }
    }

# ============================================
# Endpoints
# ============================================

@router.post("/analyze")
async def analyze_uploaded_file(
    file: UploadFile = File(...),
    current_user: Dict = Depends(get_current_user)
):
    """Upload a file and get AI-extracted metadata (name, type, parties, dates, value)"""
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"File type not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}")
    
    content = await file.read()
    
    # Save temp file for extraction
    temp_id = str(uuid.uuid4())
    temp_path = str(UPLOAD_DIR / f"temp_{temp_id}{file_ext}")
    with open(temp_path, "wb") as f:
        f.write(content)
    
    # Extract text
    extracted_text = extract_file_text(temp_path, file_ext, content)
    
    # Clean up temp file
    try:
        os.remove(temp_path)
    except:
        pass
    
    if not extracted_text or len(extracted_text.strip()) < 20:
        return {
            "success": False,
            "message": "Cannot read file content. Please enter information manually.",
            "extracted_text": extracted_text,
            "metadata": {}
        }
    
    # AI analyze
    metadata = await ai_analyze_contract(extracted_text, company_id=str(current_user["company_id"]))
    
    return {
        "success": True,
        "extracted_text": extracted_text[:5000],
        "metadata": metadata
    }


@router.post("")
async def create_contract(
    name: str = Form(""),
    contract_type: Optional[str] = Form(None),
    parties: Optional[str] = Form("[]"),
    start_date: Optional[str] = Form(None),
    end_date: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    value: Optional[str] = Form("0"),
    file: Optional[UploadFile] = File(None),
    current_user: Dict = Depends(get_current_user)
):
    """Create/upload contract with optional AI auto-fill"""
    
    # Parse parties JSON
    try:
        parties_list = json.loads(parties) if parties else []
    except:
        parties_list = []
    
    # Parse dates
    start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date() if start_date else None
    end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date() if end_date else None
    
    # Handle file upload
    file_path = None
    file_type = None
    extracted_text = None
    
    if file:
        # Validate file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        
        # Save file to temp directory
        content = await file.read()
        file_id = str(uuid.uuid4())
        file_name = f"{file_id}{file_ext}"
        file_path = str(UPLOAD_DIR / file_name)
        
        try:
            with open(file_path, "wb") as f:
                f.write(content)
        except Exception as e:
            print(f"Warning: Could not save file to disk: {e}")
            file_path = None  # Store in DB only
        
        file_type = file_ext
        
        # Extract text from uploaded file
        extracted_text = extract_file_text(file_path, file_ext, content) if file_path else None
        # Fall back to raw bytes for text files
        if not extracted_text and file_ext == ".txt":
            extracted_text = content.decode('utf-8', errors='ignore')
    
    # Insert into database
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            INSERT INTO contracts (
                company_id, uploaded_by, name, contract_type, parties,
                start_date, end_date, file_path, file_type, extracted_text, notes
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING *
        """, (
            current_user["company_id"],
            current_user["id"],
            name,
            contract_type,
            json.dumps(parties_list),
            start_date_obj,
            end_date_obj,
            file_path,
            file_type,
            extracted_text,
            notes
        ))
        
        contract = dict(cur.fetchone())
        conn.commit()
        
        # Audit log for upload
        try:
            from src.api.main import log_audit
            log_audit(str(current_user["company_id"]), str(current_user.get("id")), "upload", "contract", str(contract["id"]))
        except Exception:
            pass
        
        return contract

@router.get("")
async def list_contracts(
    contract_type: Optional[str] = None,
    status: Optional[str] = None,
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    limit: int = Query(50, le=200),
    offset: int = 0,
    current_user: Dict = Depends(get_current_user)
):
    """List contracts for company"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        query = """
            SELECT 
                c.*,
                u.full_name as uploaded_by_name,
                u.email as uploaded_by_email
            FROM contracts c
            LEFT JOIN users u ON u.id = c.uploaded_by
            WHERE c.company_id = %s AND c.status != 'deleted'
        """
        params = [current_user["company_id"]]
        
        if contract_type:
            query += " AND c.contract_type = %s"
            params.append(contract_type)
        
        if status:
            query += " AND c.status = %s"
            params.append(status)
        
        if start_date:
            query += " AND c.start_date >= %s"
            params.append(start_date)
        
        if end_date:
            query += " AND c.end_date <= %s"
            params.append(end_date)
        
        query += " ORDER BY c.created_at DESC LIMIT %s OFFSET %s"
        params.extend([limit, offset])
        
        cur.execute(query, params)
        contracts = [dict(r) for r in cur.fetchall()]
        
        return {"contracts": contracts}

@router.get("/expiring")
async def get_expiring_contracts(
    days: int = Query(30, ge=1, le=365),
    current_user: Dict = Depends(get_current_user)
):
    """Get contracts expiring within N days"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            SELECT 
                c.*,
                u.full_name as uploaded_by_name,
                (c.end_date - CURRENT_DATE) as days_until_expiry
            FROM contracts c
            LEFT JOIN users u ON u.id = c.uploaded_by
            WHERE c.company_id = %s
              AND c.end_date IS NOT NULL
              AND c.end_date > CURRENT_DATE
              AND c.end_date <= CURRENT_DATE + %s * INTERVAL '1 day'
              AND c.status = 'active'
            ORDER BY c.end_date ASC
        """, (current_user["company_id"], days))
        
        contracts = [dict(r) for r in cur.fetchall()]
        
        return {"expiring_contracts": contracts, "days": days}

@router.get("/{contract_id}")
async def get_contract(contract_id: str, current_user: Dict = Depends(get_current_user)):
    """Get contract detail"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            SELECT 
                c.*,
                u.full_name as uploaded_by_name,
                u.email as uploaded_by_email
            FROM contracts c
            LEFT JOIN users u ON u.id = c.uploaded_by
            WHERE c.id = %s AND c.company_id = %s
        """, (contract_id, current_user["company_id"]))
        
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # Audit log for view
        try:
            from src.api.main import log_audit
            log_audit(str(current_user["company_id"]), str(current_user.get("user_id")), "view", "contract", contract_id)
        except Exception:
            pass
        
        return dict(contract)

@router.put("/{contract_id}")
async def update_contract(
    contract_id: str,
    update: ContractUpdate,
    current_user: Dict = Depends(get_current_user)
):
    """Update contract metadata"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Build update query
        updates = []
        params = []
        
        if update.name is not None:
            updates.append("name = %s")
            params.append(update.name)
        
        if update.contract_type is not None:
            updates.append("contract_type = %s")
            params.append(update.contract_type)
        
        if update.parties is not None:
            updates.append("parties = %s::jsonb")
            params.append(json.dumps(update.parties))
        
        if update.start_date is not None:
            updates.append("start_date = %s")
            params.append(update.start_date)
        
        if update.end_date is not None:
            updates.append("end_date = %s")
            params.append(update.end_date)
        
        if update.status is not None:
            updates.append("status = %s")
            params.append(update.status)
        
        if update.notes is not None:
            updates.append("notes = %s")
            params.append(update.notes)
        
        if update.metadata is not None:
            updates.append("metadata = %s::jsonb")
            params.append(json.dumps(update.metadata))
        
        if not updates:
            raise HTTPException(status_code=400, detail="No updates provided")
        
        updates.append("updated_at = now()")
        
        query = f"""
            UPDATE contracts 
            SET {', '.join(updates)}
            WHERE id = %s AND company_id = %s
            RETURNING *
        """
        params.extend([contract_id, current_user["company_id"]])
        
        cur.execute(query, params)
        updated = cur.fetchone()
        
        if not updated:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        conn.commit()
        return dict(updated)

@router.delete("/{contract_id}")
async def delete_contract(
    contract_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """Soft delete contract"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            UPDATE contracts
            SET status = 'deleted', updated_at = now()
            WHERE id = %s AND company_id = %s
            RETURNING id
        """, (contract_id, current_user["company_id"]))
        
        deleted = cur.fetchone()
        if not deleted:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        conn.commit()
        
        # Audit log for delete
        try:
            from src.api.main import log_audit
            log_audit(str(current_user["company_id"]), str(current_user.get("user_id")), "delete", "contract", contract_id)
        except Exception:
            pass
        
        return {"message": "Contract deleted", "id": deleted["id"]}

@router.post("/{contract_id}/review")
async def review_contract(
    contract_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """Send contract for AI review"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get contract
        cur.execute("""
            SELECT * FROM contracts
            WHERE id = %s AND company_id = %s
        """, (contract_id, current_user["company_id"]))
        
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        contract = dict(contract)
        
        # Get text (either extracted or from file)
        text = contract.get("extracted_text")
        
        if not text:
            raise HTTPException(
                status_code=400,
                detail="No text available for review. Upload a text file or extract text from PDF/DOCX first."
            )
        
        # Call Claude for review
        result = await call_llm_for_review(text, contract.get("contract_type"), company_id=str(current_user["company_id"]))
        
        # Save review result
        cur.execute("""
            UPDATE contracts
            SET review_result = %s::jsonb, updated_at = now()
            WHERE id = %s
            RETURNING *
        """, (json.dumps(result["review"]), contract_id))
        
        updated = dict(cur.fetchone())
        conn.commit()
        
        # Update usage
        cur.execute("""
            UPDATE companies 
            SET used_quota = used_quota + 1 
            WHERE id = %s
        """, (current_user["company_id"],))
        
        cur.execute("""
            INSERT INTO usage_logs (
                company_id, user_id, endpoint, agent_type,
                input_tokens, output_tokens, status_code
            )
            VALUES (%s, %s, '/v1/contracts/review', 'review', %s, %s, 200)
        """, (
            current_user["company_id"],
            current_user["id"],
            result["tokens"]["input"],
            result["tokens"]["output"]
        ))
        
        conn.commit()
        
        return {
            "contract": updated,
            "review": result["review"],
            "tokens_used": result["tokens"]["input"] + result["tokens"]["output"]
        }

@router.post("/review-text")
async def review_contract_text(
    review: ContractReview,
    current_user: Dict = Depends(get_current_user)
):
    """Review contract from direct text input (no upload)"""
    if not review.contract_text:
        raise HTTPException(status_code=400, detail="contract_text is required")
    
    result = await call_llm_for_review(review.contract_text, company_id=str(current_user["company_id"]))
    
    # Update usage
    with get_db() as conn:
        cur = conn.cursor()
        
        cur.execute("""
            UPDATE companies 
            SET used_quota = used_quota + 1 
            WHERE id = %s
        """, (current_user["company_id"],))
        
        cur.execute("""
            INSERT INTO usage_logs (
                company_id, user_id, endpoint, agent_type,
                input_tokens, output_tokens, status_code
            )
            VALUES (%s, %s, '/v1/contracts/review-text', 'review', %s, %s, 200)
        """, (
            current_user["company_id"],
            current_user["id"],
            result["tokens"]["input"],
            result["tokens"]["output"]
        ))
        
        conn.commit()
    
    return {
        "review": result["review"],
        "tokens_used": result["tokens"]["input"] + result["tokens"]["output"]
    }


# ============================================
# Contract Review AI — NEW ENDPOINTS
# ============================================

@router.post("/{contract_id}/review-ai")
async def review_contract_ai(
    contract_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """
    AI-powered contract review using ContractReviewService
    Analyzes risks, compliance, and provides detailed recommendations
    """
    from ...services.contract_review import ContractReviewService
    
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        # Get contract
        cur.execute("""
            SELECT * FROM contracts
            WHERE id = %s AND company_id = %s AND status != 'deleted'
        """, (contract_id, current_user["company_id"]))
        
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        contract = dict(contract)
        
        # Get text
        text = contract.get("extracted_text") or contract.get("content")
        
        if not text or len(text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail="Contract has no content to analyse. Please upload the contract file."
            )
        
        # Run contract review service
        reviewer = ContractReviewService()
        
        # Parse parties
        parties = contract.get("parties")
        if parties and isinstance(parties, str):
            try:
                parties = json.loads(parties)
            except:
                parties = None
        
        review_result = reviewer.review_contract(
            contract_text=text,
            contract_name=contract.get("name", "Contract"),
            contract_type=contract.get("contract_type"),
            parties=parties
        )
        
        # Save review to database
        cur.execute("""
            UPDATE contracts
            SET review_result = %s::jsonb,
                metadata = COALESCE(metadata, '{}'::jsonb) || 
                           jsonb_build_object('last_reviewed_at', %s, 'review_score', %s),
                updated_at = NOW()
            WHERE id = %s
        """, (
            json.dumps(review_result),
            datetime.now().isoformat(),
            review_result["risk_score"],
            contract_id
        ))
        conn.commit()
        
        # Update usage
        cur.execute("UPDATE companies SET used_quota = used_quota + 1 WHERE id = %s", 
                   (current_user["company_id"],))
        conn.commit()
        
        return {
            "success": True,
            "review": review_result,
            "message": "Contract reviewed successfully"
        }


@router.get("/{contract_id}/review-ai")
async def get_contract_review_ai(
    contract_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """Get existing AI review for a contract"""
    with get_db() as conn:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        cur.execute("""
            SELECT review_result FROM contracts
            WHERE id = %s AND company_id = %s AND status != 'deleted'
        """, (contract_id, current_user["company_id"]))
        
        contract = cur.fetchone()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        review_result = contract.get("review_result")
        if not review_result:
            raise HTTPException(
                status_code=404,
                detail="No review found. Please run POST /contracts/{id}/review-ai first"
            )
        
        return {
            "success": True,
            "review": review_result
        }
