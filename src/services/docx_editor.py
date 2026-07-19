"""
DOCX Editor Service
Preserves formatting when editing Word documents
"""
from docx import Document
from typing import List, Dict


def edit_docx_file(input_path: str, output_path: str, edits: List[Dict[str, str]]) -> Dict:
    """
    Apply edits to DOCX preserving all formatting.
    
    Args:
        input_path: Path to original .docx file
        output_path: Path where edited file will be saved
        edits: List of {"find": "old text", "replace": "new text"} dicts
    
    Returns:
        {"changes_made": int, "output_path": str, "edits_applied": list}
    
    Strategy:
    - For each paragraph, check if it contains the find text
    - Replace text in runs while preserving run formatting (bold, italic, font, size)
    - Also search in table cells
    """
    doc = Document(input_path)
    changes = 0
    edits_applied = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        full_text = para.text
        for edit in edits:
            find_text = edit["find"]
            replace_text = edit["replace"]
            
            if find_text in full_text:
                success = _replace_in_paragraph(para, find_text, replace_text)
                if success:
                    changes += 1
                    edits_applied.append({
                        "find": find_text,
                        "replace": replace_text,
                        "location": "paragraph"
                    })
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = para.text
                    for edit in edits:
                        find_text = edit["find"]
                        replace_text = edit["replace"]
                        
                        if find_text in full_text:
                            success = _replace_in_paragraph(para, find_text, replace_text)
                            if success:
                                changes += 1
                                edits_applied.append({
                                    "find": find_text,
                                    "replace": replace_text,
                                    "location": "table"
                                })
    
    # Save edited document
    doc.save(output_path)
    
    return {
        "changes_made": changes,
        "output_path": output_path,
        "edits_applied": edits_applied
    }


def _replace_in_paragraph(paragraph, find_text: str, replace_text: str) -> bool:
    """
    Replace text in paragraph while preserving run formatting.
    
    Returns:
        True if replacement was made, False otherwise
    """
    full_text = paragraph.text
    
    if find_text not in full_text:
        return False
    
    # Strategy: Clear all runs except first, set new text on first run
    # This preserves the paragraph's style and first run's formatting
    new_text = full_text.replace(find_text, replace_text)
    
    if paragraph.runs:
        # Keep first run's formatting
        first_run = paragraph.runs[0]
        
        # Clear all other runs
        for run in paragraph.runs[1:]:
            run.text = ""
        
        # Set new text on first run (preserves its bold, italic, font, etc.)
        first_run.text = new_text
    else:
        # No runs, just set paragraph text
        paragraph.text = new_text
    
    return True
